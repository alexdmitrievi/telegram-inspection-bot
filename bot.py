import os
import logging
import asyncio
import json
import tempfile
import re
from datetime import datetime
from telegram import Update, ReplyKeyboardMarkup, InlineKeyboardMarkup, InlineKeyboardButton, BotCommand
from telegram.ext import (
    ApplicationBuilder, CommandHandler, MessageHandler,
    ContextTypes, ConversationHandler, filters, CallbackQueryHandler
)
from docx import Document
import nest_asyncio

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Состояния
SELECT_TEMPLATE, ASKING, CONFIRMING, BLOCK_INPUT, BLOCK_CONFIRM, BLOCK_DATE = range(6)
PROFILE_PATH = "user_profile.json"

# Вопросы для инспекции
questions = [
    "Выберите или введите наименование товара",
    "Введите массу партии в тоннах",
    "Введите количество мест",
    "Введите транспортное средство",
    "Введите дату и номер контракта/распоряжения",
    "Введите наименование отправителя",
    "Введите сопроводительные документы (инвойс и CMR)",
    "Введите дополнительные сведения",
    "Введите дату исходящего письма и инспекции"
]

# Сопоставление с переменными в Word
mapping_keys = [
    "{{TNVED_CODE}}", "{{WEIGHT}}", "{{PLACES}}", "{{VEHICLE}}", "{{CONTRACT_INFO}}",
    "{{SENDER}}", "{{DOCS}}", "{{EXTRA_INFO}}", "{{DATE}}", "{{PRODUCT_NAME}}"
]

# Справочник ТН ВЭД
product_to_tnved = {
    "лук": "0703101900", "помидор": "0702000000", "томат": "0702000000",
    "огурец": "0707009000", "перец": "0709601000", "морковь": "0706101000",
    "капуста": "0701909000", "яблоко": "0808108000", "груша": "0808209000",
    "инжир": "0804200000", "арбуз": "0807110000", "виноград": "0806101000"
}

def detect_tnved_code(name):
    name = name.lower()
    for key, code in product_to_tnved.items():
        if key in name:
            return code
    return "0808108000"

def reorder_answers(raw):
    return [
        raw[0], raw[2], raw[3], raw[4], raw[5],
        raw[6], raw[7], raw[8], raw[9], raw[1],
    ]

def save_profile(data):
    with open(PROFILE_PATH, "w", encoding="utf-8") as f:
        json.dump(dict(zip(mapping_keys, data)), f, ensure_ascii=False, indent=2)

def replace_all(doc, replacements):
    def process_paragraph(paragraph):
        full_text = ''.join(run.text for run in paragraph.runs)
        for key, val in replacements.items():
            if key in full_text:
                full_text = full_text.replace(key, val)
        for i in range(len(paragraph.runs)):
            paragraph.runs[i].text = ''
        if paragraph.runs:
            paragraph.runs[0].text = full_text

    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

def generate_statement_doc_with_date(replacements):
    template_path = "Заявление на осмотр.docx"
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)

    output_path = os.path.join(output_dir, f"Заявление_на_осмотр_{timestamp}.docx")

    doc = Document(template_path)
    replace_all(doc, replacements)
    doc.save(output_path)
    return output_path

def generate_inspection_doc_from_dict(replacements):
    template_path = "Заявка на проведение инспекции.docx"
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)

    output_path = os.path.join(output_dir, f"Заявка_на_проведение_инспекции_{timestamp}.docx")

    doc = Document(template_path)
    replace_all(doc, replacements)
    doc.save(output_path)
    return output_path

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data.clear()
    reply_markup = ReplyKeyboardMarkup([
        ["\U0001F4E6 Заявка на проведение инспекции", "\U0001F4C4 Заявление на осмотр"]
    ], resize_keyboard=True)
    await update.message.reply_text("Выберите шаблон:", reply_markup=reply_markup)
    return SELECT_TEMPLATE

async def select_template(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.lower()

    if "инспекц" in text or "заявка" in text:
        context.user_data.clear()
        context.user_data["template"] = "inspection"
        context.user_data["answers"] = []
        context.user_data["step"] = 0

        if os.path.exists(PROFILE_PATH):
            with open(PROFILE_PATH, "r", encoding="utf-8") as f:
                context.user_data["cached"] = json.load(f)
            reply_markup = ReplyKeyboardMarkup([["✅ Да", "✏ Ввести заново"]], resize_keyboard=True)
            await update.message.reply_text("🧠 Использовать данные из последней заявки?", reply_markup=reply_markup)
            return CONFIRMING
        else:
            return await prompt_product_choice(update, context)

    elif "ввести заново" in text:
        context.user_data["answers"] = []
        context.user_data["step"] = 0
        return await prompt_product_choice(update, context)

    else:
        context.user_data["template"] = "statement"
        context.user_data["blocks"] = []
        context.user_data["block_step"] = 0
        await update.message.reply_text("Введите госномер:")
        return BLOCK_INPUT

async def confirm(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip().lower()
    answers = context.user_data.get("answers", {})

    use_cache = (
        text in ["да", "\u2705 да"] and
        context.user_data.get("step") == 0 and
        "cached" in context.user_data
    )

    if use_cache:
        answers = context.user_data["cached"]
        context.user_data["answers"] = answers

    save_profile(answers)
    file = generate_inspection_doc_from_dict(answers)
    await update.message.reply_document(document=open(file, "rb"))
    return ConversationHandler.END

async def prompt_product_choice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [[InlineKeyboardButton(name.capitalize(), callback_data=name)]
                for name in list(product_to_tnved.keys())[:6]]
    reply_markup = InlineKeyboardMarkup(keyboard)

    if update.message:
        await update.message.reply_text("Выберите товар или введите вручную:", reply_markup=reply_markup)
    elif update.callback_query:
        await update.callback_query.message.reply_text("Выберите товар или введите вручную:", reply_markup=reply_markup)

    return ASKING

async def handle_inline_selection(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    return await process_step(query.message, context, query.data)

async def ask_question(update: Update, context: ContextTypes.DEFAULT_TYPE):
    return await process_step(update.message, context, update.message.text)

async def process_step(msg, context, text):
    step = context.user_data.get("step", 0)

    if not isinstance(context.user_data.get("answers"), dict):
        context.user_data["answers"] = {}

    answers = context.user_data["answers"]

    key_order = [
        "{{TNVED_CODE}}", "{{PRODUCT_NAME}}", "{{WEIGHT}}", "{{PLACES}}", "{{VEHICLE}}",
        "{{CONTRACT_INFO}}", "{{SENDER}}", "{{DOCS}}", "{{EXTRA_INFO}}", "{{DATE}}"
    ]

    if step == 0:
        product = text.strip()
        answers["{{PRODUCT_NAME}}"] = product
        answers["{{TNVED_CODE}}"] = detect_tnved_code(product)
    else:
        key = key_order[step + 1]  # сдвиг на 1, так как 0-й шаг — двойной
        answers[key] = text.strip()

    context.user_data["answers"] = answers
    context.user_data["step"] = step + 1

    if context.user_data["step"] < len(questions):
        await msg.reply_text(questions[context.user_data["step"]])
        return ASKING
    else:
        summary = "\n".join([
            f"{questions[i]}: {answers.get(key_order[i + 1 if i > 0 else 1], '—')}"
            for i in range(len(questions))
        ])
        await msg.reply_text(
            f"Проверьте введённые данные:\n\n{summary}\n\nОтправить документы? (да/нет)",
            reply_markup=ReplyKeyboardMarkup([
                ["🔄 Перезапустить", "да", "нет"]
            ], resize_keyboard=True)
        )
        return CONFIRMING

async def block_input(update: Update, context: ContextTypes.DEFAULT_TYPE):
    step = context.user_data.get("block_step", 0)
    if step == 0:
        context.user_data["v"] = update.message.text.strip()
        context.user_data["block_step"] = 1
        await update.message.reply_text("Введите документы:")
        return BLOCK_INPUT
    elif step == 1:
        context.user_data["d"] = update.message.text.strip()
        context.user_data["block_step"] = 2
        await update.message.reply_text("Введите товар:")
        return BLOCK_INPUT
    else:
        product = update.message.text.strip()
        context.user_data["blocks"].append(f"г/н {context.user_data['v']} по {context.user_data['d']}, товар: {product}")
        context.user_data["block_step"] = 0
        reply_markup = ReplyKeyboardMarkup([["➕ Да", "✅ Нет"]], resize_keyboard=True)
        await update.message.reply_text("Добавить ещё?", reply_markup=reply_markup)
        return BLOCK_CONFIRM

async def confirm_blocks(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if "да" in update.message.text.lower():
        await update.message.reply_text("Введите госномер:")
        return BLOCK_INPUT
    else:
        await update.message.reply_text("Введите дату осмотра:")
        return BLOCK_DATE

async def set_block_date(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["date"] = update.message.text.strip()
    blocks = context.user_data.get("blocks", [])
    replacements = {
        "{{BLOCKS}}": "\n".join(blocks),
        "{{DATE}}": context.user_data.get("date", "")
    }
    file = generate_statement_doc_with_date(replacements)
    await update.message.reply_document(document=open(file, "rb"))
    return ConversationHandler.END

async def run():
    token = os.getenv("BOT_TOKEN")
    app = ApplicationBuilder().token(token).build()

    conv = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            SELECT_TEMPLATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, select_template)],
            CONFIRMING: [MessageHandler(filters.TEXT & ~filters.COMMAND, confirm)],
            ASKING: [
                CallbackQueryHandler(handle_inline_selection),
                MessageHandler(filters.TEXT & ~filters.COMMAND, ask_question)
            ],
            BLOCK_INPUT: [MessageHandler(filters.TEXT & ~filters.COMMAND, block_input)],
            BLOCK_CONFIRM: [MessageHandler(filters.TEXT & ~filters.COMMAND, confirm_blocks)],
            BLOCK_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, set_block_date)],
        },
        fallbacks=[CommandHandler("start", start)],
    )

    app.add_handler(conv)
    await app.bot.set_my_commands([BotCommand("start", "Начать заполнение заявки")])
    await app.run_polling()

if __name__ == '__main__':
    nest_asyncio.apply()
    asyncio.run(run())



